"""
Code syntax highlighting renderer.
Sử dụng Pygments để tạo syntax highlighting cho code blocks trong DOCX.
"""

import logging

from pygments import lex
from pygments.lexers import get_lexer_by_name, TextLexer, guess_lexer
from pygments.token import Token

from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from markdocx.styles import (
    FONT_CODE,
    FONT_SIZE_CODE_BLOCK,
    COLOR_CODE_BLOCK_BG,
    COLOR_CODE_BLOCK_BORDER,
    get_syntax_color,
    set_paragraph_shading,
    set_paragraph_borders,
)

logger = logging.getLogger(__name__)

# Language aliases for common names
LANGUAGE_ALIASES = {
    "py": "python",
    "js": "javascript",
    "ts": "typescript",
    "rb": "ruby",
    "cs": "csharp",
    "c++": "cpp",
    "c#": "csharp",
    "sh": "bash",
    "shell": "bash",
    "yml": "yaml",
    "md": "markdown",
    "tex": "latex",
    "rs": "rust",
    "kt": "kotlin",
    "m": "objectivec",
    "dockerfile": "docker",
    "plaintext": "text",
    "plain": "text",
    "txt": "text",
    "": "text",
}


def get_lexer(language):
    """Get a Pygments lexer for the given language."""
    lang = language.strip().lower() if language else ""
    lang = LANGUAGE_ALIASES.get(lang, lang)

    try:
        return get_lexer_by_name(lang, stripall=False)
    except Exception:
        try:
            return get_lexer_by_name("text", stripall=False)
        except Exception:
            return TextLexer()


def tokenize_code(code, language):
    """
    Tokenize source code using Pygments.

    Args:
        code: Source code string
        language: Programming language name

    Returns:
        List of (token_type, token_value) tuples
    """
    lexer = get_lexer(language)
    tokens = list(lex(code, lexer))
    return tokens


def add_code_block_to_doc(doc, code, language=""):
    """
    Add a syntax-highlighted code block to the document.

    Args:
        doc: python-docx Document
        code: Source code string
        language: Programming language name

    Returns:
        The created paragraph
    """
    # Remove trailing newline if present
    code = code.rstrip("\n")
    if not code:
        return None

    tokens = tokenize_code(code, language)

    # Add language label if specified
    lang_display = language.strip() if language else ""
    if lang_display:
        label_para = doc.add_paragraph()
        label_para.paragraph_format.space_after = Pt(0)
        label_para.paragraph_format.space_before = Pt(8)
        label_run = label_para.add_run(f"  {lang_display}")
        label_run.font.name = FONT_CODE
        label_run.font.size = Pt(8)
        label_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        label_run.font.italic = True
        # Set East Asian font for the run
        rPr = label_run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_CODE)

    # Create the code paragraph
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2) if lang_display else Pt(8)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.0

    # Style: gray background + borders
    set_paragraph_shading(para, COLOR_CODE_BLOCK_BG)
    set_paragraph_borders(para, color=COLOR_CODE_BLOCK_BORDER, size="4", space="6")

    # Add padding via indentation
    para.paragraph_format.left_indent = Pt(10)
    para.paragraph_format.right_indent = Pt(10)

    # Render each token with syntax coloring
    for token_type, token_value in tokens:
        if not token_value:
            continue

        color = get_syntax_color(token_type)

        # Split token value by newlines
        parts = token_value.split("\n")
        for idx, part in enumerate(parts):
            if part:
                run = para.add_run(part)
                run.font.name = FONT_CODE
                run.font.size = FONT_SIZE_CODE_BLOCK
                if color:
                    run.font.color.rgb = color
                # East Asian font
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), FONT_CODE)

            # Add line break between lines (not after last line)
            if idx < len(parts) - 1:
                br_run = para.add_run()
                br_run.add_break()

    return para
