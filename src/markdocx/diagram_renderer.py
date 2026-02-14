"""
Diagram, chart, graph, matrix, and workflow renderer.
Renders special fenced code blocks (```matrix, ```chart, ```graph, ```workflow)
as images and embeds them into the DOCX document.

Uses matplotlib for all rendering, with networkx for graph layouts.
"""

import json
import logging
import re
from io import BytesIO

import matplotlib
matplotlib.use("Agg")  # Non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ── Defaults ────────────────────────────────────────────────────────
DEFAULT_DPI = 200
DEFAULT_MAX_WIDTH = Inches(5.5)

# ── Color palette ───────────────────────────────────────────────────
PALETTE = [
    "#4285F4", "#EA4335", "#FBBC04", "#34A853",
    "#FF6D01", "#46BDC6", "#7B61FF", "#F538A0",
    "#185ABC", "#B31412", "#E37400", "#0D652D",
]


# ════════════════════════════════════════════════════════════════════
# Public API — called by docx_builder
# ════════════════════════════════════════════════════════════════════

DIAGRAM_LANGUAGES = {"matrix", "chart", "graph", "workflow"}


def is_diagram_language(language: str) -> bool:
    """Check if a fenced code block language is a diagram type."""
    return language.strip().lower() in DIAGRAM_LANGUAGES


def render_diagram_to_doc(doc, code: str, language: str, max_width=DEFAULT_MAX_WIDTH):
    """
    Render a diagram from a fenced code block and add it to the DOCX document.

    Args:
        doc: python-docx Document
        code: The content of the fenced code block
        language: The language identifier (matrix, chart, graph, workflow)
        max_width: Maximum image width in the document
    """
    lang = language.strip().lower()

    try:
        if lang == "matrix":
            img_bytes, caption = _render_matrix(code)
        elif lang == "chart":
            img_bytes, caption = _render_chart(code)
        elif lang == "graph":
            img_bytes, caption = _render_graph(code)
        elif lang == "workflow":
            img_bytes, caption = _render_workflow(code)
        else:
            logger.warning(f"Unknown diagram language: {language}")
            return

        if img_bytes is None:
            # Fallback: show raw content as text
            para = doc.add_paragraph()
            run = para.add_run(f"[Failed to render {lang} diagram]")
            run.italic = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            return

        # Add the image to the document
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(BytesIO(img_bytes), width=max_width)

        # Add caption if provided
        if caption:
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_para.add_run(caption)
            cap_run.font.size = Pt(9)
            cap_run.italic = True
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    except Exception as e:
        logger.error(f"Failed to render {lang} diagram: {e}")
        para = doc.add_paragraph()
        run = para.add_run(f"[Error rendering {lang} diagram: {e}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


# ════════════════════════════════════════════════════════════════════
# Matrix Renderer
# ════════════════════════════════════════════════════════════════════

def _parse_matrix_block(code: str):
    """
    Parse matrix block content.

    Supported formats:
    1. Simple space/comma-separated values:
       1 2 3
       4 5 6
       7 8 9

    2. JSON format:
       {"name": "A", "data": [[1,2],[3,4]], "caption": "Matrix A"}

    Returns:
        (name, data, caption) tuple
    """
    code = code.strip()

    # Try JSON first
    if code.startswith("{"):
        try:
            obj = json.loads(code)
            name = obj.get("name", "")
            data = obj.get("data", [])
            caption = obj.get("caption", "")
            return name, data, caption
        except json.JSONDecodeError:
            pass

    # Parse simple text format
    lines = [line.strip() for line in code.strip().split("\n") if line.strip()]

    name = ""
    caption = ""
    data_lines = []

    for line in lines:
        # Check for name: or caption: directives
        if line.lower().startswith("name:"):
            name = line[5:].strip()
            continue
        if line.lower().startswith("caption:"):
            caption = line[8:].strip()
            continue
        # Parse data row
        # Support both comma and space separation
        if "," in line:
            row = [_try_float(x.strip()) for x in line.split(",") if x.strip()]
        else:
            row = [_try_float(x) for x in line.split() if x]
        if row:
            data_lines.append(row)

    return name, data_lines, caption


def _try_float(s):
    """Try to convert string to float, return string if not possible."""
    try:
        f = float(s)
        if f == int(f):
            return int(f)
        return f
    except (ValueError, TypeError):
        return s


def _render_matrix(code: str):
    """Render a matrix as a styled image."""
    name, data, caption = _parse_matrix_block(code)

    if not data:
        return None, None

    rows = len(data)
    cols = max(len(row) for row in data)

    # Pad rows to same length
    for row in data:
        while len(row) < cols:
            row.append("")

    fig, ax = plt.subplots(figsize=(max(2, cols * 0.8 + 1.5), max(1.5, rows * 0.5 + 1)))
    ax.set_xlim(-0.5, cols - 0.5)
    ax.set_ylim(-0.5, rows - 0.5)
    ax.set_aspect("equal")
    ax.axis("off")

    # Draw bracket lines
    bracket_x_left = -0.65
    bracket_x_right = cols - 0.35
    bracket_y_top = -0.6
    bracket_y_bottom = rows - 0.4

    # Left bracket
    ax.plot([bracket_x_left + 0.15, bracket_x_left, bracket_x_left, bracket_x_left + 0.15],
            [bracket_y_top, bracket_y_top, bracket_y_bottom, bracket_y_bottom],
            color="black", linewidth=2, solid_capstyle="round")

    # Right bracket
    ax.plot([bracket_x_right - 0.15, bracket_x_right, bracket_x_right, bracket_x_right - 0.15],
            [bracket_y_top, bracket_y_top, bracket_y_bottom, bracket_y_bottom],
            color="black", linewidth=2, solid_capstyle="round")

    # Draw cell values
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            ax.text(j, i, str(val), ha="center", va="center",
                    fontsize=20, fontfamily="serif")

    # Add matrix name
    if name:
        ax.text(bracket_x_left - 0.4, (rows - 1) / 2, f"{name} =",
                ha="right", va="center", fontsize=22, fontfamily="serif",
                fontstyle="italic")

    plt.tight_layout(pad=0.5)

    img_bytes = _fig_to_bytes(fig)
    return img_bytes, caption or None


# ════════════════════════════════════════════════════════════════════
# Chart Renderer
# ════════════════════════════════════════════════════════════════════

def _parse_chart_block(code: str):
    """
    Parse chart block content.

    Supported format (JSON):
    {
        "type": "bar" | "line" | "pie" | "scatter",
        "title": "Chart Title",
        "xlabel": "X Axis",
        "ylabel": "Y Axis",
        "data": {
            "labels": ["A", "B", "C"],
            "datasets": [
                {"label": "Series 1", "values": [10, 20, 30]},
                {"label": "Series 2", "values": [15, 25, 35]}
            ]
        },
        "caption": "Figure 1: Description"
    }

    Or simple key-value format:
    type: bar
    title: My Chart
    labels: A, B, C, D
    Series 1: 10, 20, 30, 40
    Series 2: 15, 25, 35, 45
    caption: Figure 1
    """
    code = code.strip()

    # Try JSON
    if code.startswith("{"):
        try:
            return json.loads(code)
        except json.JSONDecodeError:
            pass

    # Parse simple format
    result = {
        "type": "bar",
        "title": "",
        "xlabel": "",
        "ylabel": "",
        "caption": "",
        "data": {"labels": [], "datasets": []}
    }

    lines = [line.strip() for line in code.split("\n") if line.strip()]
    for line in lines:
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        key = key.strip().lower()
        value = value.strip()

        if key == "type":
            result["type"] = value.lower()
        elif key == "title":
            result["title"] = value
        elif key == "xlabel":
            result["xlabel"] = value
        elif key == "ylabel":
            result["ylabel"] = value
        elif key == "caption":
            result["caption"] = value
        elif key == "labels":
            result["data"]["labels"] = [x.strip() for x in value.split(",")]
        else:
            # Treat as dataset
            values = []
            for v in value.split(","):
                v = v.strip()
                try:
                    values.append(float(v))
                except ValueError:
                    values.append(0)
            result["data"]["datasets"].append({
                "label": key.strip(),
                "values": values
            })

    # Re-parse with original case for dataset labels
    for line in [l.strip() for l in code.split("\n") if l.strip()]:
        if ":" not in line:
            continue
        key, _, value = line.partition(":")
        k_lower = key.strip().lower()
        if k_lower not in ("type", "title", "xlabel", "ylabel", "caption", "labels"):
            # Find matching dataset and fix label
            for ds in result["data"]["datasets"]:
                if ds["label"] == k_lower:
                    ds["label"] = key.strip()
                    break

    return result


def _render_chart(code: str):
    """Render a chart as an image."""
    spec = _parse_chart_block(code)

    chart_type = spec.get("type", "bar").lower()
    title = spec.get("title", "")
    xlabel = spec.get("xlabel", "")
    ylabel = spec.get("ylabel", "")
    caption = spec.get("caption", "")
    data = spec.get("data", {})
    labels = data.get("labels", [])
    datasets = data.get("datasets", [])

    if not datasets:
        return None, None

    fig, ax = plt.subplots(figsize=(8, 5))

    if chart_type == "pie":
        # Pie chart — use first dataset
        values = datasets[0].get("values", [])
        if not labels:
            labels = [f"Item {i+1}" for i in range(len(values))]
        colors = PALETTE[:len(values)]
        wedges, texts, autotexts = ax.pie(
            values, labels=labels, autopct="%1.1f%%",
            colors=colors, startangle=90,
            textprops={"fontsize": 14}
        )
        for autotext in autotexts:
            autotext.set_fontsize(13)
            autotext.set_color("white")
            autotext.set_fontweight("bold")

    elif chart_type == "scatter":
        for idx, ds in enumerate(datasets):
            values = ds.get("values", [])
            x_vals = list(range(len(values))) if not labels else list(range(len(values)))
            color = PALETTE[idx % len(PALETTE)]
            ax.scatter(x_vals, values, label=ds.get("label", f"Series {idx+1}"),
                       color=color, s=60, zorder=3)
        if labels:
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontsize=13)
        ax.legend(fontsize=13)
        ax.grid(True, alpha=0.3)

    elif chart_type == "line":
        for idx, ds in enumerate(datasets):
            values = ds.get("values", [])
            x_vals = list(range(len(values)))
            color = PALETTE[idx % len(PALETTE)]
            ax.plot(x_vals, values, label=ds.get("label", f"Series {idx+1}"),
                    color=color, marker="o", linewidth=2, markersize=5)
        if labels:
            ax.set_xticks(range(len(labels)))
            ax.set_xticklabels(labels, fontsize=13)
        ax.legend(fontsize=13)
        ax.grid(True, alpha=0.3)

    else:  # bar (default)
        n_datasets = len(datasets)
        n_labels = max(len(ds.get("values", [])) for ds in datasets)
        if not labels:
            labels = [f"Item {i+1}" for i in range(n_labels)]

        if n_datasets == 1:
            values = datasets[0].get("values", [])
            colors = PALETTE[:len(values)]
            bars = ax.bar(range(len(values)), values, color=colors, width=0.6)
            # Add value labels on top of bars
            for bar_item in bars:
                height = bar_item.get_height()
                ax.text(bar_item.get_x() + bar_item.get_width() / 2., height,
                        f"{height:g}", ha="center", va="bottom", fontsize=13)
        else:
            bar_width = 0.8 / n_datasets
            for idx, ds in enumerate(datasets):
                values = ds.get("values", [])
                x_positions = [x + idx * bar_width - (n_datasets - 1) * bar_width / 2
                               for x in range(len(values))]
                color = PALETTE[idx % len(PALETTE)]
                ax.bar(x_positions, values, bar_width,
                       label=ds.get("label", f"Series {idx+1}"), color=color)
            ax.legend(fontsize=13)

        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, fontsize=13)
        ax.grid(True, axis="y", alpha=0.3)

    if title:
        ax.set_title(title, fontsize=18, fontweight="bold", pad=12)
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=15)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=15)

    plt.tight_layout()
    img_bytes = _fig_to_bytes(fig)
    return img_bytes, caption or None


# ════════════════════════════════════════════════════════════════════
# Graph Renderer
# ════════════════════════════════════════════════════════════════════

def _parse_graph_block(code: str):
    """
    Parse graph block content.

    Supported format (JSON):
    {
        "directed": true,
        "nodes": ["A", "B", "C"],
        "edges": [
            {"from": "A", "to": "B", "label": "5"},
            {"from": "B", "to": "C", "label": "3"}
        ],
        "title": "Graph Title",
        "caption": "Figure: Description"
    }

    Or simple edge-list format:
    directed: true
    title: My Graph
    A -> B: 5
    B -> C: 3
    A -- C: 7
    caption: Figure 1
    """
    code = code.strip()

    # Try JSON
    if code.startswith("{"):
        try:
            return json.loads(code)
        except json.JSONDecodeError:
            pass

    # Parse simple format
    result = {
        "directed": False,
        "nodes": [],
        "edges": [],
        "title": "",
        "caption": "",
    }

    nodes_set = set()
    lines = [line.strip() for line in code.split("\n") if line.strip()]

    for line in lines:
        lower = line.lower()

        if lower.startswith("directed:"):
            val = line.split(":", 1)[1].strip().lower()
            result["directed"] = val in ("true", "yes", "1")
            continue
        if lower.startswith("title:"):
            result["title"] = line.split(":", 1)[1].strip()
            continue
        if lower.startswith("caption:"):
            result["caption"] = line.split(":", 1)[1].strip()
            continue
        if lower.startswith("nodes:"):
            node_list = line.split(":", 1)[1].strip()
            result["nodes"] = [n.strip() for n in node_list.split(",") if n.strip()]
            nodes_set.update(result["nodes"])
            continue

        # Parse edge: "A -> B: label" or "A -- B: label" or "A -> B"
        edge_match = re.match(
            r"^\s*(.+?)\s*(->|-->|--|—)\s*(.+?)(?:\s*:\s*(.+?))?\s*$", line
        )
        if edge_match:
            src = edge_match.group(1).strip()
            arrow = edge_match.group(2).strip()
            dst = edge_match.group(3).strip()
            label = edge_match.group(4).strip() if edge_match.group(4) else ""

            nodes_set.add(src)
            nodes_set.add(dst)

            is_directed = arrow in ("->", "-->")
            if is_directed:
                result["directed"] = True

            result["edges"].append({
                "from": src,
                "to": dst,
                "label": label
            })

    # Merge discovered nodes
    if not result["nodes"]:
        result["nodes"] = list(nodes_set)

    return result


def _render_graph(code: str):
    """Render a graph/network diagram as an image."""
    try:
        import networkx as nx
    except ImportError:
        logger.warning("networkx not installed. Install with: pip install networkx")
        return None, None

    spec = _parse_graph_block(code)

    nodes = spec.get("nodes", [])
    edges = spec.get("edges", [])
    directed = spec.get("directed", False)
    title = spec.get("title", "")
    caption = spec.get("caption", "")

    if not nodes and not edges:
        return None, None

    # Build networkx graph
    G = nx.DiGraph() if directed else nx.Graph()
    G.add_nodes_from(nodes)

    edge_labels = {}
    for e in edges:
        src = e.get("from", "")
        dst = e.get("to", "")
        label = e.get("label", "")
        G.add_edge(src, dst)
        if label:
            edge_labels[(src, dst)] = label

    # Layout
    if len(G.nodes) <= 2:
        pos = nx.spring_layout(G, seed=42, k=2)
    else:
        try:
            pos = nx.kamada_kawai_layout(G)
        except Exception:
            pos = nx.spring_layout(G, seed=42, k=2)

    fig, ax = plt.subplots(figsize=(8, 6))

    # Draw nodes
    node_colors = [PALETTE[i % len(PALETTE)] for i in range(len(G.nodes))]
    nx.draw_networkx_nodes(G, pos, ax=ax, node_color=node_colors,
                           node_size=1200, alpha=0.9)

    # Draw labels
    nx.draw_networkx_labels(G, pos, ax=ax, font_size=15,
                            font_weight="bold", font_color="white")

    # Draw edges
    if directed:
        nx.draw_networkx_edges(
            G, pos, ax=ax, edge_color="#555555",
            arrows=True, arrowsize=20, arrowstyle="-|>",
            connectionstyle="arc3,rad=0.1", width=1.5
        )
    else:
        nx.draw_networkx_edges(G, pos, ax=ax, edge_color="#555555", width=1.5)

    # Draw edge labels
    if edge_labels:
        nx.draw_networkx_edge_labels(G, pos, edge_labels, ax=ax,
                                     font_size=14, font_color="#CC0000",
                                     bbox=dict(boxstyle="round,pad=0.2",
                                               facecolor="white",
                                               edgecolor="none", alpha=0.8))

    if title:
        ax.set_title(title, fontsize=18, fontweight="bold", pad=12)

    ax.axis("off")
    plt.tight_layout()
    img_bytes = _fig_to_bytes(fig)
    return img_bytes, caption or None


# ════════════════════════════════════════════════════════════════════
# Workflow Renderer
# ════════════════════════════════════════════════════════════════════

def _parse_workflow_block(code: str):
    """
    Parse workflow block content.

    Supported format (JSON):
    {
        "title": "Workflow Title",
        "direction": "vertical" | "horizontal",
        "steps": [
            {"text": "Start", "type": "start"},
            {"text": "Process Data", "type": "process"},
            {"text": "Valid?", "type": "decision"},
            {"text": "Save Results", "type": "process"},
            {"text": "End", "type": "end"}
        ],
        "caption": "Figure: Workflow description"
    }

    Or simple format:
    title: My Workflow
    direction: vertical
    [Start]
    (Process Data)
    {Valid?}
    (Save Results)
    [End]
    caption: Figure 1

    Notation:
    - [text]  → start/end (rounded box)
    - (text)  → process (rectangle)
    - {text}  → decision (diamond)
    - <text>  → input/output (parallelogram)
    """
    code = code.strip()

    # Try JSON
    if code.startswith("{"):
        try:
            return json.loads(code)
        except json.JSONDecodeError:
            pass

    # Parse simple format
    result = {
        "title": "",
        "direction": "vertical",
        "steps": [],
        "caption": "",
    }

    lines = [line.strip() for line in code.split("\n") if line.strip()]

    for line in lines:
        lower = line.lower()

        if lower.startswith("title:"):
            result["title"] = line.split(":", 1)[1].strip()
            continue
        if lower.startswith("direction:"):
            result["direction"] = line.split(":", 1)[1].strip().lower()
            continue
        if lower.startswith("caption:"):
            result["caption"] = line.split(":", 1)[1].strip()
            continue

        # Parse step notation
        step = _parse_step(line)
        if step:
            result["steps"].append(step)

    return result


def _parse_step(line: str):
    """Parse a single workflow step line."""
    line = line.strip()
    if not line:
        return None

    # [text] → start/end
    m = re.match(r"^\[(.+)\]$", line)
    if m:
        return {"text": m.group(1).strip(), "type": "terminal"}

    # (text) → process
    m = re.match(r"^\((.+)\)$", line)
    if m:
        return {"text": m.group(1).strip(), "type": "process"}

    # {text} → decision
    m = re.match(r"^\{(.+)\}$", line)
    if m:
        return {"text": m.group(1).strip(), "type": "decision"}

    # <text> → input/output
    m = re.match(r"^<(.+)>$", line)
    if m:
        return {"text": m.group(1).strip(), "type": "io"}

    # Plain text → process
    return {"text": line, "type": "process"}


def _render_workflow(code: str):
    """Render a workflow/flowchart as an image."""
    spec = _parse_workflow_block(code)

    steps = spec.get("steps", [])
    title = spec.get("title", "")
    direction = spec.get("direction", "vertical")
    caption = spec.get("caption", "")

    if not steps:
        return None, None

    n = len(steps)
    is_vertical = direction.startswith("v")

    # ── Measure text to compute adaptive box sizes ──────────────
    def _wrap_text(text, max_chars):
        """Wrap text into multiple lines if too long."""
        if len(text) <= max_chars:
            return text
        words = text.split()
        lines = []
        current = ""
        for word in words:
            test = f"{current} {word}".strip() if current else word
            if len(test) <= max_chars:
                current = test
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
        return "\n".join(lines) if lines else text

    # Find longest text to size boxes appropriately
    max_text_len = max(len(s.get("text", "")) for s in steps)

    if is_vertical:
        # Vertical: boxes can be wide
        max_chars_per_line = 30
        box_w = max(3.0, min(6.0, max_text_len * 0.18 + 1.0))
        box_h = 1.1
        fontsize_normal = 18
        fontsize_decision = 16
        spacing_y = 2.2  # center-to-center vertical distance

        fig_w = max(6, box_w + 3.0)
        fig_h = max(4, n * spacing_y + 2.0)
    else:
        # Horizontal with multi-row zigzag layout
        # Determine how many steps per row (max 4 per row for readability)
        cols_per_row = min(4, n)  # at most 4 boxes per row
        if n > 4:
            cols_per_row = (n + 1) // 2  # split roughly in half
            cols_per_row = min(cols_per_row, 4)  # cap at 4

        max_chars_per_line = 16
        box_w = max(3.2, min(5.0, max_text_len * 0.18 + 1.5))
        box_h = 1.6
        fontsize_normal = 24
        fontsize_decision = 22
        spacing_x = box_w + 1.5   # horizontal gap between boxes
        spacing_y = 3.5           # vertical gap between rows

        num_rows = (n + cols_per_row - 1) // cols_per_row
        fig_w = max(10, cols_per_row * spacing_x + 2.5)
        fig_h = max(5, num_rows * spacing_y + 3.0)

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")

    # ── Calculate positions ─────────────────────────────────────
    positions = []
    if is_vertical:
        for i in range(n):
            x = fig_w / 2
            y = fig_h - 1.5 - i * spacing_y
            positions.append((x, y))
    else:
        # Zigzag: row 0 left→right, row 1 right→left, row 2 left→right, ...
        for i in range(n):
            row = i // cols_per_row
            col = i % cols_per_row
            # Reverse direction on odd rows
            if row % 2 == 1:
                col = cols_per_row - 1 - col
            x = 1.5 + col * spacing_x
            y = fig_h - 1.8 - row * spacing_y
            positions.append((x, y))

    # Adjust axes limits
    all_x = [p[0] for p in positions]
    all_y = [p[1] for p in positions]
    pad_x = box_w / 2 + 1.5
    pad_y = box_h + 1.5
    ax.set_xlim(min(all_x) - pad_x, max(all_x) + pad_x)
    ax.set_ylim(min(all_y) - pad_y, max(all_y) + pad_y)

    # ── Color scheme ────────────────────────────────────────────
    type_colors = {
        "terminal": ("#34A853", "white"),
        "start": ("#34A853", "white"),
        "end": ("#EA4335", "white"),
        "process": ("#4285F4", "white"),
        "decision": ("#FBBC04", "black"),
        "io": ("#7B61FF", "white"),
    }

    # ── Draw steps ──────────────────────────────────────────────
    for i, step in enumerate(steps):
        x, y = positions[i]
        step_type = step.get("type", "process")
        raw_text = step.get("text", "")
        display_text = _wrap_text(raw_text, max_chars_per_line)
        bg_color, text_color = type_colors.get(step_type, ("#4285F4", "white"))

        # Count lines for height adjustment
        n_lines = display_text.count("\n") + 1
        effective_h = max(box_h, box_h + (n_lines - 1) * 0.3)

        if step_type == "decision":
            diamond_w = box_w * 0.55
            diamond_h = effective_h * 0.9
            diamond = plt.Polygon([
                (x, y + diamond_h),
                (x + diamond_w, y),
                (x, y - diamond_h),
                (x - diamond_w, y)
            ], closed=True, facecolor=bg_color, edgecolor="#333333",
                linewidth=1.5, zorder=3)
            ax.add_patch(diamond)
            ax.text(x, y, display_text, ha="center", va="center",
                    fontsize=fontsize_decision, fontweight="bold",
                    color=text_color, zorder=4, linespacing=1.2)

        elif step_type in ("terminal", "start", "end"):
            patch = FancyBboxPatch(
                (x - box_w / 2, y - effective_h / 2), box_w, effective_h,
                boxstyle="round,pad=0.15", facecolor=bg_color,
                edgecolor="#333333", linewidth=1.5, zorder=3
            )
            ax.add_patch(patch)
            ax.text(x, y, display_text, ha="center", va="center",
                    fontsize=fontsize_normal, fontweight="bold",
                    color=text_color, zorder=4, linespacing=1.2)

        elif step_type == "io":
            skew = 0.25
            pgram = plt.Polygon([
                (x - box_w / 2 + skew, y + effective_h / 2),
                (x + box_w / 2 + skew, y + effective_h / 2),
                (x + box_w / 2 - skew, y - effective_h / 2),
                (x - box_w / 2 - skew, y - effective_h / 2)
            ], closed=True, facecolor=bg_color, edgecolor="#333333",
                linewidth=1.5, zorder=3)
            ax.add_patch(pgram)
            ax.text(x, y, display_text, ha="center", va="center",
                    fontsize=fontsize_normal, fontweight="bold",
                    color=text_color, zorder=4, linespacing=1.2)

        else:  # process
            patch = FancyBboxPatch(
                (x - box_w / 2, y - effective_h / 2), box_w, effective_h,
                boxstyle="square,pad=0.1", facecolor=bg_color,
                edgecolor="#333333", linewidth=1.5, zorder=3
            )
            ax.add_patch(patch)
            ax.text(x, y, display_text, ha="center", va="center",
                    fontsize=fontsize_normal, fontweight="bold",
                    color=text_color, zorder=4, linespacing=1.2)

    # ── Draw arrows between steps ───────────────────────────────
    for i in range(n - 1):
        x1, y1 = positions[i]
        x2, y2 = positions[i + 1]

        cur_type = steps[i].get("type", "process")
        next_type = steps[i + 1].get("type", "process")

        cur_text = _wrap_text(steps[i].get("text", ""), max_chars_per_line)
        next_text = _wrap_text(steps[i + 1].get("text", ""), max_chars_per_line)
        cur_h = max(box_h, box_h + (cur_text.count("\n")) * 0.3)
        next_h = max(box_h, box_h + (next_text.count("\n")) * 0.3)

        same_row = abs(y1 - y2) < 0.1  # same horizontal row

        if is_vertical:
            if cur_type == "decision":
                start_y = y1 - cur_h * 0.9 - 0.08
            else:
                start_y = y1 - cur_h / 2 - 0.08
            if next_type == "decision":
                end_y = y2 + next_h * 0.9 + 0.08
            else:
                end_y = y2 + next_h / 2 + 0.08
            ax.annotate("", xy=(x2, end_y), xytext=(x1, start_y),
                        arrowprops=dict(arrowstyle="-|>", color="#555555",
                                        lw=2, mutation_scale=15),
                        zorder=2)
        elif same_row:
            # Horizontal arrow on same row
            if x2 > x1:
                # Left to right
                start_x = x1 + box_w / 2 + 0.08
                end_x = x2 - box_w / 2 - 0.08
            else:
                # Right to left
                start_x = x1 - box_w / 2 - 0.08
                end_x = x2 + box_w / 2 + 0.08
            if cur_type == "decision":
                start_x = x1 + (box_w * 0.55 + 0.08) * (1 if x2 > x1 else -1)
            if next_type == "decision":
                end_x = x2 - (box_w * 0.55 + 0.08) * (1 if x2 > x1 else -1)
            ax.annotate("", xy=(end_x, y2), xytext=(start_x, y1),
                        arrowprops=dict(arrowstyle="-|>", color="#555555",
                                        lw=2, mutation_scale=15),
                        zorder=2)
        else:
            # Vertical arrow between rows (row transition)
            if cur_type == "decision":
                start_y = y1 - cur_h * 0.9 - 0.08
            else:
                start_y = y1 - cur_h / 2 - 0.08
            if next_type == "decision":
                end_y = y2 + next_h * 0.9 + 0.08
            else:
                end_y = y2 + next_h / 2 + 0.08
            ax.annotate("", xy=(x2, end_y), xytext=(x1, start_y),
                        arrowprops=dict(arrowstyle="-|>", color="#555555",
                                        lw=2, mutation_scale=15,
                                        connectionstyle="arc3,rad=0"),
                        zorder=2)

    if title:
        ax.set_title(title, fontsize=28, fontweight="bold", pad=20)

    plt.tight_layout()
    img_bytes = _fig_to_bytes(fig)
    return img_bytes, caption or None


# ════════════════════════════════════════════════════════════════════
# Utility
# ════════════════════════════════════════════════════════════════════

def _fig_to_bytes(fig, dpi=DEFAULT_DPI):
    """Convert a matplotlib figure to PNG bytes."""
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.read()
