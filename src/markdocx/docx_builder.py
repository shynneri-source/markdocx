"""
DOCX document builder.
Chuyển đổi token stream từ markdown-it-py thành tài liệu DOCX.
Xử lý tất cả các phần tử: heading, paragraph, list, table, code, math, image, link...
"""

import os
import logging
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from markdocx.styles import (
    FONT_BODY,
    FONT_CODE,
    FONT_SIZE_BODY,
    FONT_SIZE_CODE,
    COLOR_BODY_TEXT,
    COLOR_LINK,
    COLOR_INLINE_CODE_BG,
    COLOR_BLOCKQUOTE_TEXT,
    LIST_INDENT,
    set_run_shading,
    set_blockquote_style,
    add_horizontal_rule,
    set_table_header_shading,
    setup_document_styles,
)
from markdocx.math_renderer import latex_to_omml, latex_to_omml_para
from markdocx.code_renderer import add_code_block_to_doc
from markdocx.diagram_renderer import is_diagram_language, render_diagram_to_doc

logger = logging.getLogger(__name__)


class DocxBuilder:
    """Builds a DOCX document from markdown-it-py tokens."""

    def __init__(self, base_dir=".", image_max_width=Inches(5.5)):
        """
        Args:
            base_dir: Base directory for resolving relative image paths
            image_max_width: Maximum width for images in the document
        """
        self.doc = Document()
        self.base_dir = os.path.abspath(base_dir)
        self.image_max_width = image_max_width
        self._blockquote_depth = 0
        self._list_depth = 0
        self._ordered_counters = {}  # depth -> counter
        self._shape_id_counter = 0

        setup_document_styles(self.doc)

    def _next_shape_id(self):
        self._shape_id_counter += 1
        return self._shape_id_counter

    # ════════════════════════════════════════════════════════════════
    # Main entry point
    # ════════════════════════════════════════════════════════════════

    def build(self, tokens, output_path):
        """
        Build a DOCX document from tokens and save it.

        Args:
            tokens: List of markdown-it-py tokens
            output_path: Path to save the DOCX file
        """
        self._process_tokens(tokens)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"Saved: {output_path}")

    # ════════════════════════════════════════════════════════════════
    # Token stream processor
    # ════════════════════════════════════════════════════════════════

    def _process_tokens(self, tokens):
        """Walk the token list and dispatch to handlers."""
        i = 0
        while i < len(tokens):
            token = tokens[i]
            t = token.type

            if t == "heading_open":
                i = self._handle_heading(tokens, i)
            elif t == "paragraph_open":
                i = self._handle_paragraph(tokens, i)
            elif t in ("fence", "code_block"):
                self._handle_code_block(token)
                i += 1
            elif t in ("math_block", "math_block_eqno"):
                self._handle_math_block(token)
                i += 1
            elif t == "bullet_list_open":
                i = self._handle_list(tokens, i, ordered=False)
            elif t == "ordered_list_open":
                i = self._handle_list(tokens, i, ordered=True)
            elif t == "blockquote_open":
                i = self._handle_blockquote(tokens, i)
            elif t == "table_open":
                i = self._handle_table(tokens, i)
            elif t == "hr":
                add_horizontal_rule(self.doc)
                i += 1
            elif t == "html_block":
                # Skip raw HTML blocks
                i += 1
            elif t == "front_matter":
                # Skip YAML front matter
                i += 1
            elif t == "footnote_block_open":
                i = self._handle_footnote_block(tokens, i)
            else:
                # Skip unknown tokens
                i += 1

    # ════════════════════════════════════════════════════════════════
    # Block handlers
    # ════════════════════════════════════════════════════════════════

    def _handle_heading(self, tokens, i):
        """Handle heading_open → inline → heading_close."""
        level = int(tokens[i].tag[1])  # h1→1, h2→2, ...
        i += 1

        if i < len(tokens) and tokens[i].type == "inline":
            inline_token = tokens[i]
            i += 1
        else:
            inline_token = None

        # Skip heading_close
        if i < len(tokens) and tokens[i].type == "heading_close":
            i += 1

        heading = self.doc.add_heading("", level=min(level, 6))
        # Clear the default run
        heading.clear()

        if inline_token and inline_token.children:
            self._process_inline(heading, inline_token.children)

        # Apply blockquote styling if inside blockquote
        if self._blockquote_depth > 0:
            set_blockquote_style(heading, self._blockquote_depth)

        return i

    def _handle_paragraph(self, tokens, i):
        """Handle paragraph_open → inline → paragraph_close."""
        i += 1  # Skip paragraph_open

        inline_token = None
        if i < len(tokens) and tokens[i].type == "inline":
            inline_token = tokens[i]
            i += 1

        # Skip paragraph_close
        if i < len(tokens) and tokens[i].type == "paragraph_close":
            i += 1

        # Check if this paragraph is a list item
        if self._list_depth > 0:
            para = self.doc.add_paragraph()
            depth = self._list_depth - 1
            para.paragraph_format.left_indent = LIST_INDENT * (depth + 1)
        else:
            para = self.doc.add_paragraph()

        if inline_token and inline_token.children:
            self._process_inline(para, inline_token.children)

        # Apply blockquote styling if inside blockquote
        if self._blockquote_depth > 0:
            set_blockquote_style(para, self._blockquote_depth)
            for run in para.runs:
                run.font.color.rgb = COLOR_BLOCKQUOTE_TEXT

        return i

    def _handle_code_block(self, token):
        """Handle fence or code_block token."""
        language = ""
        if token.info:
            language = token.info.strip().split()[0]
        code = token.content

        # Check if this is a diagram block (matrix, chart, graph, workflow)
        if is_diagram_language(language):
            render_diagram_to_doc(self.doc, code, language, self.image_max_width)
        else:
            add_code_block_to_doc(self.doc, code, language)

    def _handle_math_block(self, token):
        """Handle display math ($$...$$) block using native OMML."""
        latex_str = token.content.strip()
        if not latex_str:
            return

        omml_para = latex_to_omml_para(latex_str)
        if omml_para is not None:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para._element.append(omml_para)
        else:
            # Fallback: show raw LaTeX in code font
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"[{latex_str}]")
            run.font.name = FONT_CODE
            run.font.size = FONT_SIZE_CODE
            run.italic = True

    def _handle_list(self, tokens, i, ordered=False):
        """Handle bullet_list or ordered_list (with nesting)."""
        self._list_depth += 1
        depth = self._list_depth - 1

        if ordered:
            self._ordered_counters[depth] = 0

        i += 1  # Skip list_open

        while i < len(tokens):
            token = tokens[i]

            if token.type in ("bullet_list_close", "ordered_list_close"):
                i += 1
                break

            if token.type == "list_item_open":
                if ordered:
                    self._ordered_counters[depth] = \
                        self._ordered_counters.get(depth, 0) + 1
                i += 1

                # Process content within list item
                while i < len(tokens) and tokens[i].type != "list_item_close":
                    if tokens[i].type == "paragraph_open":
                        i += 1  # Skip paragraph_open

                        inline_token = None
                        if i < len(tokens) and tokens[i].type == "inline":
                            inline_token = tokens[i]
                            i += 1

                        if i < len(tokens) and tokens[i].type == "paragraph_close":
                            i += 1

                        # Create list item paragraph
                        para = self.doc.add_paragraph()
                        para.paragraph_format.left_indent = LIST_INDENT * (depth + 1)
                        para.paragraph_format.space_before = Pt(1)
                        para.paragraph_format.space_after = Pt(1)

                        # Add bullet/number prefix
                        bullets = ["•", "◦", "▪", "▹", "•", "◦"]
                        if ordered:
                            num = self._ordered_counters.get(depth, 1)
                            prefix_run = para.add_run(f"{num}. ")
                            prefix_run.font.name = FONT_BODY
                            prefix_run.font.size = FONT_SIZE_BODY
                            prefix_run.bold = True
                        else:
                            bullet_char = bullets[min(depth, len(bullets) - 1)]
                            prefix_run = para.add_run(f"{bullet_char}  ")
                            prefix_run.font.name = FONT_BODY
                            prefix_run.font.size = FONT_SIZE_BODY

                        # Add content
                        if inline_token and inline_token.children:
                            self._process_inline(para, inline_token.children)

                        # Apply blockquote if needed
                        if self._blockquote_depth > 0:
                            set_blockquote_style(para, self._blockquote_depth)

                    elif tokens[i].type == "bullet_list_open":
                        i = self._handle_list(tokens, i, ordered=False)
                    elif tokens[i].type == "ordered_list_open":
                        i = self._handle_list(tokens, i, ordered=True)
                    elif tokens[i].type in ("fence", "code_block"):
                        self._handle_code_block(tokens[i])
                        i += 1
                    elif tokens[i].type in ("math_block", "math_block_eqno"):
                        self._handle_math_block(tokens[i])
                        i += 1
                    elif tokens[i].type == "blockquote_open":
                        i = self._handle_blockquote(tokens, i)
                    else:
                        i += 1

                # Skip list_item_close
                if i < len(tokens) and tokens[i].type == "list_item_close":
                    i += 1
            else:
                i += 1

        self._list_depth -= 1
        if ordered and depth in self._ordered_counters:
            del self._ordered_counters[depth]

        return i

    def _handle_blockquote(self, tokens, i):
        """Handle blockquote by collecting inner tokens and processing them."""
        self._blockquote_depth += 1
        i += 1  # Skip blockquote_open
        depth = 1

        inner_tokens = []
        while i < len(tokens):
            if tokens[i].type == "blockquote_open":
                depth += 1
            elif tokens[i].type == "blockquote_close":
                depth -= 1
                if depth == 0:
                    i += 1
                    break
            inner_tokens.append(tokens[i])
            i += 1

        self._process_tokens(inner_tokens)
        self._blockquote_depth -= 1
        return i

    def _handle_table(self, tokens, i):
        """Parse table tokens and create a Word table."""
        i += 1  # Skip table_open

        rows_data = []
        current_row = []
        is_header_section = False

        while i < len(tokens):
            token = tokens[i]

            if token.type == "table_close":
                i += 1
                break
            elif token.type == "thead_open":
                is_header_section = True
                i += 1
            elif token.type == "thead_close":
                is_header_section = False
                i += 1
            elif token.type in ("tbody_open", "tbody_close"):
                i += 1
            elif token.type == "tr_open":
                current_row = []
                i += 1
            elif token.type == "tr_close":
                rows_data.append((current_row[:], is_header_section))
                i += 1
            elif token.type in ("th_open", "td_open"):
                # Get alignment from token attrs
                align = None
                if token.attrs:
                    style_val = token.attrGet("style") if hasattr(token, 'attrGet') else None
                    if style_val and "text-align:center" in str(style_val):
                        align = WD_ALIGN_PARAGRAPH.CENTER
                    elif style_val and "text-align:right" in str(style_val):
                        align = WD_ALIGN_PARAGRAPH.RIGHT
                i += 1
                # Next should be inline
                inline_tok = None
                if i < len(tokens) and tokens[i].type == "inline":
                    inline_tok = tokens[i]
                    i += 1
                current_row.append((inline_tok, is_header_section, align))
                # Skip th_close / td_close
                if i < len(tokens) and tokens[i].type in ("th_close", "td_close"):
                    i += 1
            else:
                i += 1

        # Build the table
        if not rows_data:
            return i

        num_cols = max(len(row) for row, _ in rows_data) if rows_data else 0
        num_rows = len(rows_data)

        if num_cols == 0 or num_rows == 0:
            return i

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = 1  # Center

        for row_idx, (row_cells, is_hdr) in enumerate(rows_data):
            for col_idx, cell_data in enumerate(row_cells):
                if col_idx >= num_cols:
                    break
                inline_tok, is_header, align = cell_data
                cell = table.rows[row_idx].cells[col_idx]
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)

                if align:
                    para.alignment = align

                if inline_tok and inline_tok.children:
                    self._process_inline(para, inline_tok.children)

                if is_header or is_hdr:
                    set_table_header_shading(cell)
                    for run in para.runs:
                        run.bold = True

        return i

    def _handle_footnote_block(self, tokens, i):
        """Handle footnote block."""
        i += 1  # Skip footnote_block_open

        # Add separator
        add_horizontal_rule(self.doc)
        title = self.doc.add_paragraph()
        run = title.add_run("Chú thích")
        run.bold = True
        run.font.size = Pt(10)

        depth = 1
        while i < len(tokens) and depth > 0:
            if tokens[i].type == "footnote_block_open":
                depth += 1
            elif tokens[i].type == "footnote_block_close":
                depth -= 1
                if depth == 0:
                    i += 1
                    break
            elif tokens[i].type == "footnote_open":
                i += 1
                continue
            elif tokens[i].type == "footnote_close":
                i += 1
                continue
            elif tokens[i].type == "paragraph_open":
                i += 1
                if i < len(tokens) and tokens[i].type == "inline":
                    inline_tok = tokens[i]
                    i += 1
                    para = self.doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.3)
                    if inline_tok.children:
                        self._process_inline(para, inline_tok.children)
                    for run in para.runs:
                        run.font.size = Pt(9)
                if i < len(tokens) and tokens[i].type == "paragraph_close":
                    i += 1
                continue

            i += 1

        return i

    # ════════════════════════════════════════════════════════════════
    # Inline content processor
    # ════════════════════════════════════════════════════════════════

    def _process_inline(self, paragraph, children):
        """
        Process inline tokens and add formatted runs to a paragraph.
        Handles: text, bold, italic, strikethrough, code, math, links, images.
        """
        if not children:
            return

        bold = False
        italic = False
        strikethrough = False
        link_href = None

        idx = 0
        while idx < len(children):
            child = children[idx]
            ct = child.type

            # ── Formatting toggles ───────────────────────────
            if ct == "strong_open":
                bold = True
                idx += 1
                continue
            elif ct == "strong_close":
                bold = False
                idx += 1
                continue
            elif ct == "em_open":
                italic = True
                idx += 1
                continue
            elif ct == "em_close":
                italic = False
                idx += 1
                continue
            elif ct == "s_open":
                strikethrough = True
                idx += 1
                continue
            elif ct == "s_close":
                strikethrough = False
                idx += 1
                continue

            # ── Links ────────────────────────────────────────
            elif ct == "link_open":
                href = child.attrGet("href") if hasattr(child, 'attrGet') else None
                if not href and child.attrs:
                    href = child.attrs.get("href", "")
                link_href = href or ""
                idx += 1
                continue
            elif ct == "link_close":
                link_href = None
                idx += 1
                continue

            # ── Plain text ───────────────────────────────────
            elif ct == "text":
                text = child.content
                if link_href:
                    self._add_hyperlink(paragraph, link_href, text,
                                       bold=bold, italic=italic)
                else:
                    run = paragraph.add_run(text)
                    self._apply_run_format(run, bold, italic, strikethrough)
                idx += 1
                continue

            # ── Inline code ──────────────────────────────────
            elif ct == "code_inline":
                run = paragraph.add_run(child.content)
                run.font.name = FONT_CODE
                run.font.size = FONT_SIZE_CODE
                set_run_shading(run, COLOR_INLINE_CODE_BG)
                # East Asian font
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:eastAsia"), FONT_CODE)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                idx += 1
                continue

            # ── Inline math ──────────────────────────────────
            elif ct == "math_inline":
                latex_str = child.content.strip()
                omml = latex_to_omml(latex_str)
                if omml is not None:
                    paragraph._element.append(omml)
                else:
                    # Fallback: show as styled text
                    run = paragraph.add_run(f"${latex_str}$")
                    run.font.name = FONT_CODE
                    run.font.size = FONT_SIZE_CODE
                    run.italic = True
                idx += 1
                continue

            # ── Inline math (display in inline context) ──────
            elif ct == "math_inline_double":
                latex_str = child.content.strip()
                omml = latex_to_omml(latex_str)
                if omml is not None:
                    paragraph._element.append(omml)
                else:
                    run = paragraph.add_run(f"$${latex_str}$$")
                    run.font.name = FONT_CODE
                    run.italic = True
                idx += 1
                continue

            # ── Image ────────────────────────────────────────
            elif ct == "image":
                self._handle_image(paragraph, child)
                idx += 1
                continue

            # ── Line breaks ──────────────────────────────────
            elif ct == "softbreak":
                run = paragraph.add_run("\n")
                idx += 1
                continue
            elif ct == "hardbreak":
                run = paragraph.add_run()
                run.add_break()
                idx += 1
                continue

            # ── HTML inline (skip) ───────────────────────────
            elif ct == "html_inline":
                # Try to handle <br> tags
                content = child.content.strip().lower()
                if content in ("<br>", "<br/>", "<br />"):
                    run = paragraph.add_run()
                    run.add_break()
                idx += 1
                continue

            # ── Footnote reference ───────────────────────────
            elif ct == "footnote_ref":
                ref_id = child.meta.get("id", "?") if hasattr(child, "meta") and child.meta else "?"
                run = paragraph.add_run(f"[{int(ref_id) + 1}]")
                run.font.size = Pt(8)
                run.font.superscript = True
                run.font.color.rgb = COLOR_LINK
                idx += 1
                continue

            else:
                # Unknown inline token - skip
                logger.debug(f"Unknown inline token: {ct}")
                idx += 1
                continue

    def _apply_run_format(self, run, bold=False, italic=False, strikethrough=False):
        """Apply formatting to a run."""
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if strikethrough:
            run.font.strike = True

    # ════════════════════════════════════════════════════════════════
    # Hyperlinks
    # ════════════════════════════════════════════════════════════════

    def _add_hyperlink(self, paragraph, url, text, bold=False, italic=False):
        """Add a clickable hyperlink to a paragraph."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)

            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")

            # Color
            color_el = OxmlElement("w:color")
            color_el.set(qn("w:val"), "0563C1")
            rPr.append(color_el)

            # Underline
            u_el = OxmlElement("w:u")
            u_el.set(qn("w:val"), "single")
            rPr.append(u_el)

            # Font
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), FONT_BODY)
            rFonts.set(qn("w:hAnsi"), FONT_BODY)
            rFonts.set(qn("w:eastAsia"), FONT_BODY)
            rPr.append(rFonts)

            # Bold / Italic
            if bold:
                b_el = OxmlElement("w:b")
                rPr.append(b_el)
            if italic:
                i_el = OxmlElement("w:i")
                rPr.append(i_el)

            new_run.append(rPr)

            text_el = OxmlElement("w:t")
            text_el.set(qn("xml:space"), "preserve")
            text_el.text = text
            new_run.append(text_el)

            hyperlink.append(new_run)
            paragraph._element.append(hyperlink)

        except Exception as e:
            logger.warning(f"Failed to create hyperlink for {url}: {e}")
            run = paragraph.add_run(text)
            run.font.color.rgb = COLOR_LINK
            run.underline = True
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    # ════════════════════════════════════════════════════════════════
    # Images
    # ════════════════════════════════════════════════════════════════

    def _handle_image(self, paragraph, token):
        """Handle an image token."""
        src = token.attrGet("src") if hasattr(token, 'attrGet') else None
        if not src and token.attrs:
            src = token.attrs.get("src", "")
        alt = ""
        if token.children:
            alt = "".join(c.content for c in token.children if c.type == "text")
        elif hasattr(token, 'attrGet'):
            alt = token.attrGet("alt") or ""

        if not src:
            run = paragraph.add_run(f"[Image: {alt}]")
            run.italic = True
            return

        # Resolve image path
        image_path = src
        if not os.path.isabs(image_path):
            image_path = os.path.join(self.base_dir, image_path)

        if os.path.exists(image_path):
            try:
                run = paragraph.add_run()
                run.add_picture(image_path, width=self.image_max_width)

                # Add alt text as caption if available
                if alt:
                    cap_para = self.doc.add_paragraph()
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_para.add_run(alt)
                    cap_run.font.size = Pt(9)
                    cap_run.italic = True
                    cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            except Exception as e:
                logger.warning(f"Failed to add image {src}: {e}")
                run = paragraph.add_run(f"[Image: {alt or src}]")
                run.italic = True
        else:
            logger.warning(f"Image not found: {image_path}")
            run = paragraph.add_run(f"[Image not found: {src}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
